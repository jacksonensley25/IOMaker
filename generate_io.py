#!/usr/bin/env python3
"""
generate_io.py

Now supports two standardized brief formats:
- Sponsored Search (PLA)
- Display

Select the field map via:
  - CLI: --campaign pla  OR  --campaign display
  - Streamlit app: radio selector

Still supports setting {{service_description}} to "PLA MS"/"PLA SS".
"""

import re
import argparse
from pathlib import Path
import datetime as dt
from typing import Dict, Tuple, Optional

from openpyxl import load_workbook
from docx import Document

# ---------- Paths ----------
BRIEF_PATH = Path("sample.xlsx")
IO_TEMPLATE_PATH = Path("sample.docx")
OUTPUT_PATH = Path("output_io.docx")
SHEET_NAME: Optional[str] = None

# ---------- Formatters ----------
def fmt_strip(v): return "" if v is None else str(v).strip()

def fmt_date(v):
    if isinstance(v, dt.datetime): return v.date().isoformat()
    if isinstance(v, dt.date):     return v.isoformat()
    s = fmt_strip(v)
    if not s: return s
    for patt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%d-%b-%Y", "%b %d, %Y"):
        try: return dt.datetime.strptime(s, patt).date().isoformat()
        except ValueError: pass
    return s

def fmt_currency(v):
    s = fmt_strip(v)
    if not s: return s
    try:
        num = float(s.replace(",", "").replace("$", ""))
        return f"${num:,.0f}" if abs(num - round(num)) < 1e-9 else f"${num:,.2f}"
    except Exception:
        return s

FORMATTERS = {"strip": fmt_strip, "date": fmt_date, "currency": fmt_currency}

# ---------- Field Maps ----------
# Each entry: label -> (cell, placeholder, formatter)

FIELD_MAP_PLA: Dict[str, Tuple[str, str, str]] = {
    "Campaign name": ("D13", "{{campaign_name}}", "strip"),
    "Client name":   ("D4",  "{{client_name}}",   "strip"),

    "Commercial contact name":  ("D32", "{{commercial_contact_name}}", "strip"),
    "Commercial contact role":  ("D33", "{{commercial_contact_role}}", "strip"),
    "Commercial contact email": ("D34", "{{commercial_contact_email}}", "strip"),

    "Financial contact name":  ("D38", "{{financial_contact_name}}", "strip"),
    "Financial contact role":  ("D39", "{{financial_contact_role}}", "strip"),
    "Financial contact email": ("D40", "{{financial_contact_email}}", "strip"),

    # service_description is injected (radio / --service)
    "Start date": ("D16", "{{start_date}}", "date"),
    "End date":   ("D17", "{{end_date}}",   "date"),

    "Campaign budget": ("D20", "{{campaign_budget}}", "currency"),
    "Total budget":    ("D20", "{{total_budget}}",    "currency"),
}

# TODO: update these cells to the real Display brief layout
FIELD_MAP_DISPLAY: Dict[str, Tuple[str, str, str]] = {
    "Campaign name": ("D13", "{{campaign_name}}", "strip"),
    "Client name":   ("D4",  "{{client_name}}",   "strip"),

    "Commercial contact name":  ("D33", "{{commercial_contact_name}}", "strip"),
    "Commercial contact role":  ("D34", "{{commercial_contact_role}}", "strip"),
    "Commercial contact email": ("D35", "{{commercial_contact_email}}", "strip"),

    "Financial contact name":  ("D39", "{{financial_contact_name}}", "strip"),
    "Financial contact role":  ("D40", "{{financial_contact_role}}", "strip"),
    "Financial contact email": ("D41", "{{financial_contact_email}}", "strip"),

    "Start date": ("D15", "{{start_date}}", "date"),
    "End date":   ("D16", "{{end_date}}",   "date"),

    "Campaign budget": ("D19", "{{campaign_budget}}", "currency"),
    "Total budget":    ("D19", "{{total_budget}}",    "currency"),
}

FIELD_MAPS = {
    "pla": FIELD_MAP_PLA,
    "display": FIELD_MAP_DISPLAY,
}

REQUIRED_PLACEHOLDERS = {
    "{{campaign_name}}", "{{client_name}}",
    "{{start_date}}", "{{end_date}}", "{{total_budget}}",
}

# ---------- Core helpers ----------
def read_brief_values(xlsx_path: Path, field_map: Dict[str, Tuple[str, str, str]]) -> Dict[str, str]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb[SHEET_NAME] if SHEET_NAME else wb.active
    repl = {}
    for _, (cell, placeholder, fmt_key) in field_map.items():
        raw = ws[cell].value
        repl[placeholder] = FORMATTERS.get(fmt_key, fmt_strip)(raw)
    return repl

def _collect_token_spans(text: str, replacements: dict):
    """Return list of (start, end, replacement_text) for each placeholder occurrence."""
    spans = []
    for ph, val in replacements.items():
        if not ph:
            continue
        start = 0
        while True:
            idx = text.find(ph, start)
            if idx == -1:
                break
            spans.append((idx, idx + len(ph), val))
            start = idx + len(ph)
    # process from right to left so index math stays valid
    spans.sort(key=lambda x: x[0], reverse=True)
    return spans

def _apply_spans_to_paragraph_preserve_runs(paragraph, spans):
    """Apply (start, end, replacement) spans to a paragraph without resetting runs."""
    if not spans:
        return
    runs = paragraph.runs
    texts = [r.text for r in runs]

    # cumulative boundaries for original text
    cum = [0]
    for t in texts:
        cum.append(cum[-1] + len(t))

    def locate(pos: int):
        """Map absolute char index -> (run_index, offset) using ORIGINAL boundaries."""
        # pos is in [0, cum[-1]]
        for i in range(len(texts)):
            if cum[i] <= pos <= cum[i + 1]:
                return i, pos - cum[i]
        # fallback: end of last run
        return len(texts) - 1, len(texts[-1])

    # apply from right to left
    for start, end, repl in spans:
        si, so = locate(start)
        ei, eo = locate(end)
        if si == ei:
            r = runs[si]
            r.text = r.text[:so] + repl + r.text[eo:]
        else:
            # keep tail in the end run with its original formatting
            tail = runs[ei].text[eo:]
            runs[ei].text = tail
            # clear all full middle runs
            for j in range(si + 1, ei):
                runs[j].text = ""
            # put left part + replacement in the start run
            runs[si].text = runs[si].text[:so] + repl

def replace_in_paragraph_preserve(paragraph, replacements: dict):
    """Replace placeholders while preserving runs (styling, images)."""
    if not paragraph.runs:
        return
    full = "".join(r.text for r in paragraph.runs)
    if not full:
        return
    spans = _collect_token_spans(full, replacements)
    _apply_spans_to_paragraph_preserve_runs(paragraph, spans)

def replace_in_cell_preserve(cell, replacements: dict):
    for p in cell.paragraphs:
        replace_in_paragraph_preserve(p, replacements)
    for t in cell.tables:
        for r in t.rows:
            for c in r.cells:
                replace_in_cell_preserve(c, replacements)

def replace_everywhere(doc, replacements: dict):
    # body
    for p in doc.paragraphs:
        replace_in_paragraph_preserve(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell_preserve(cell, replacements)
    # headers/footers
    for section in doc.sections:
        for hdrftr in (section.header, section.footer):
            for p in hdrftr.paragraphs:
                replace_in_paragraph_preserve(p, replacements)
            for table in hdrftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_cell_preserve(cell, replacements)


def validate_required(repl: Dict[str, str]) -> Optional[str]:
    missing = [ph for ph in REQUIRED_PLACEHOLDERS if not repl.get(ph)]
    return f"Missing required value(s) for: {', '.join(missing)}" if missing else None

# ----------------------- Service description (campaign + service) ------------
SERVICE_DESC_MAP = {
    ("pla", "managed"):  "PLA MS",
    ("pla", "self"):     "PLA SS",
    ("display", "managed"): "Display MS",
    ("display", "self"):    "Display SS",
}

def _normalize_service(s: str) -> str:
    s = (s or "").strip().lower()
    # accept "managed", "managed service", etc.
    return "managed" if s.startswith("managed") else "self"

def _normalize_campaign(c: str) -> str:
    c = (c or "").strip().lower()
    # accept "pla", "sponsored search (pla)", etc.
    if c.startswith("pla") or "sponsored" in c:
        return "pla"
    return "display"

def compute_service_description(campaign: str, service: str) -> str:
    ck = _normalize_campaign(campaign)
    sk = _normalize_service(service)
    return SERVICE_DESC_MAP.get((ck, sk), "PLA MS")

def service_to_description(choice: str) -> str:
    return SERVICE_MAP.get((choice or "").strip().lower(), "PLA MS")

def compute_billing_unit_and_rate(campaign: str) -> tuple[str, str]:
    """
    Returns (billing_unit, rate_string) based on campaign type.
    - Display  => ("CPM", "$20")
    - PLA      => ("CPC", "$0.80")
    """
    c = (campaign or "").strip().lower()
    if c.startswith("pla") or "sponsored" in c:
        return "CPC", "$0.80"
    return "CPM", "$20"

# ---------- Main ----------
def main(service: Optional[str] = "managed", campaign: str = "pla"):
    # 1) Pick map
    fmap = FIELD_MAPS.get(campaign.lower(), FIELD_MAP_PLA)

    # 2) Excel → replacements
    values = read_brief_values(BRIEF_PATH, fmap)

    # 3) Inject service description
    values["{{service_description}}"] = compute_service_description(campaign, service)

    # 3b) Inject billing unit + rate from campaign type
    bu, rt = compute_billing_unit_and_rate(campaign)
    values["{{billing_unit}}"] = bu
    values["{{rate}}"] = rt


    # 4) Validate + generate
    err = validate_required(values)
    if err: print(f"⚠️  {err}")

    doc = Document(IO_TEMPLATE_PATH)
    replace_everywhere(doc, values)
    doc.save(OUTPUT_PATH)
    print(f"✅ Insertion Order generated: {OUTPUT_PATH.resolve()}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate IO from brief (PLA/Display) with service description.")
    parser.add_argument("--campaign", choices=["pla", "display"], default="pla", help="Brief type / field map.")
    parser.add_argument("--service", choices=["managed", "self", "Managed Service", "Self Service", "managed service", "self service"], default="managed", help="Service type for {{service_description}}.")
    args = parser.parse_args()
    main(service=args.service, campaign=args.campaign)
