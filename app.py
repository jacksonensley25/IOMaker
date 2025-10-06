import io
import re
import datetime as dt
import tempfile
from pathlib import Path

import streamlit as st
from docx import Document

# your generator module (must include: FIELD_MAPS, read_brief_values, replace_everywhere,
# validate_required, compute_service_description, compute_billing_unit_and_rate,
# IO_TEMPLATE_PATH, fmt_currency)
import generate_io as gio


# ---------- Page config ----------
st.set_page_config(page_title="Brief → Insertion Order", page_icon="🧾", layout="centered")
st.title("BJ's IO Maker")

# Subtle button styling (roomier, rounded)
st.markdown(
    """
    <style>
    .stButton > button {
        padding: 0.9rem 1.2rem;
        border-radius: 12px;
        font-weight: 600;
        border: 1px solid #eaeaea;
    }
    .stButton > button:hover {
        border-color: #c9c9c9;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

with st.expander("How it works", expanded=False):
    st.markdown(
        """
Fill **Section 1** and click **Add another brief** to reveal **Section 2**. Do the same in
Section 2 to reveal **Section 3**. Or click **Generate IO** at any time.

- **Brief Type** selects the field map (Sponsored Search/PLA vs Display).
- **Service Type** + **Brief Type** sets `{{service_description}}` (PLA MS/SS, Display MS/SS).
- **Billing unit & rate** auto-set by Brief Type (PLA → CPC $0.80, Display → CPM $20).
- Campaigns 2 & 3 write to suffixed placeholders like `{{campaign_name_2}}`, `{{campaign_name_3}}`.
        """
    )

# ---------- Session flags for showing sections ----------
if "show_c2" not in st.session_state:
    st.session_state.show_c2 = False
if "show_c3" not in st.session_state:
    st.session_state.show_c3 = False


# ---------- Helpers ----------
def suffix_placeholders(repl: dict, suffix: str) -> dict:
    if not suffix:
        return dict(repl)
    out = {}
    for k, v in repl.items():
        if k.startswith("{{") and k.endswith("}}"):
            inner = k[2:-2]
            out[f"{{{{{inner}{suffix}}}}}"] = v
        else:
            out[k + suffix] = v
    return out

def _currency_to_float(x):
    if x is None:
        return 0.0
    if isinstance(x, (int, float)):
        return float(x)
    s = str(x).strip().replace("$", "").replace(",", "")
    try:
        return float(s) if s else 0.0
    except ValueError:
        return 0.0

def _remove_placeholders_with_suffixes(doc, suffixes):
    """Remove any {{placeholder_<suffix>}} without resetting runs or touching images."""
    if not suffixes:
        return
    suffix_alt = "|".join(re.escape(s) for s in suffixes)  # e.g. "_2|_3"
    pattern = re.compile(r"\{\{[^{}]+(?:" + suffix_alt + r")\}\}")

    def _clean_paragraph(p):
        runs = p.runs
        if not runs:
            return
        full = "".join(r.text for r in runs)
        if not full:
            return
        spans = [(m.start(), m.end()) for m in pattern.finditer(full)]
        if not spans:
            return

        texts = [r.text for r in runs]
        cum = [0]
        for t in texts:
            cum.append(cum[-1] + len(t))

        def locate(pos: int):
            for i in range(len(texts)):
                if cum[i] <= pos <= cum[i + 1]:
                    return i, pos - cum[i]
            return len(texts) - 1, len(texts[-1])

        for start, end in sorted(spans, key=lambda x: x[0], reverse=True):
            si, so = locate(start)
            ei, eo = locate(end)
            if si == ei:
                runs[si].text = runs[si].text[:so] + runs[si].text[eo:]
            else:
                tail = runs[ei].text[eo:]
                runs[ei].text = tail
                for j in range(si + 1, ei):
                    runs[j].text = ""
                runs[si].text = runs[si].text[:so]

    def _clean_cell(cell):
        for p in cell.paragraphs:
            _clean_paragraph(p)
        for t in cell.tables:
            for r in t.rows:
                for c in r.cells:
                    _clean_cell(c)

    # body
    for p in doc.paragraphs:
        _clean_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _clean_cell(cell)

    # headers/footers
    for section in doc.sections:
        for hdrftr in (section.header, section.footer):
            for p in hdrftr.paragraphs:
                _clean_paragraph(p)
            for table in hdrftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _clean_cell(cell)


def campaign_section(title: str, key_prefix: str):
    st.subheader(title)
    colA, colB = st.columns(2)
    with colA:
        brief_type = st.selectbox(
            "Brief Type",
            options=["Sponsored Search (PLA)", "Display"],
            index=0,
            key=f"{key_prefix}_brief_type",
        )
    with colB:
        service_choice = st.selectbox(
            "Service Type",
            options=["Managed Service", "Self Service"],
            index=0,
            key=f"{key_prefix}_service",
        )
    brief_file = st.file_uploader("Upload brief (.xlsx)", type=["xlsx"], key=f"{key_prefix}_uploader")
    return brief_type, service_choice, brief_file


# ---------- UI: Sections ----------
brief_type_1, service_choice_1, file_1 = campaign_section("Campaign #1", "c1")

file_2 = file_3 = None
brief_type_2 = service_choice_2 = None
brief_type_3 = service_choice_3 = None

if st.session_state.show_c2:
    st.divider()
    brief_type_2, service_choice_2, file_2 = campaign_section("Campaign #2", "c2")

if st.session_state.show_c3:
    st.divider()
    brief_type_3, service_choice_3, file_3 = campaign_section("Campaign #3", "c3")


# ---------- BOTTOM CONTROL BAR (always at the bottom) ----------
st.divider()
left, right = st.columns(2, gap="large")

with left:
    # Add logic: reveal Section 2, then Section 3; disable when both are visible
    if not st.session_state.show_c2:
        add_label = "Add another brief"
    elif not st.session_state.show_c3:
        add_label = "Add another brief (final)"
    else:
        add_label = "Add another brief (max 3 reached)"
    add_disabled = st.session_state.show_c2 and st.session_state.show_c3
    add_clicked = st.button(add_label, key="btn_add", use_container_width=True, disabled=add_disabled)

with right:
    gen_btn = st.button("Generate IO", type="primary", key="btn_gen", use_container_width=True)

# Handle add click (after layout is drawn so buttons stay bottom)
if add_clicked:
    if not st.session_state.show_c2:
        st.session_state.show_c2 = True
    elif not st.session_state.show_c3:
        st.session_state.show_c3 = True
    st.rerun()


# ---------- Generate IO ----------
if gen_btn:
    # Require at least section 1 file
    if not file_1:
        st.error("Please upload the brief for Section 1 (Campaign 1).")
        st.stop()

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir = Path(tmpdir)
            template_path = Path(gio.IO_TEMPLATE_PATH)
            if not template_path.exists():
                st.error("Default template not found at 'sample.docx'. Please add it.")
                st.stop()

            master_repl = {}

            def process(uploaded_file, brief_type, service_choice, suffix: str):
                if not uploaded_file:
                    return False
                # Save brief
                brief_path = tmpdir / f"uploaded_brief{suffix or '_1'}.xlsx"
                brief_path.write_bytes(uploaded_file.getvalue())

                # Map field set
                campaign_key = "pla" if (brief_type and "PLA" in brief_type) else "display"
                field_map = gio.FIELD_MAPS[campaign_key]

                # Excel → values
                repl = gio.read_brief_values(brief_path, field_map)

                # Service desc
                repl["{{service_description}}"] = gio.compute_service_description(campaign_key, service_choice)

                # Billing unit + rate by type
                bu, rt = gio.compute_billing_unit_and_rate(campaign_key)
                repl["{{billing_unit}}"] = bu
                repl["{{rate}}"] = rt

                # Suffix & merge
                repl = suffix_placeholders(repl, suffix)
                master_repl.update(repl)
                return True

            used_1 = process(file_1, brief_type_1, service_choice_1, "")
            used_2 = used_3 = False
            if st.session_state.show_c2:
                used_2 = process(file_2, brief_type_2, service_choice_2, "_2")
            if st.session_state.show_c3:
                used_3 = process(file_3, brief_type_3, service_choice_3, "_3")

            # Grand total across all campaign budgets
            grand_total = 0.0
            for key in ("{{campaign_budget}}", "{{campaign_budget_2}}", "{{campaign_budget_3}}"):
                grand_total += _currency_to_float(master_repl.get(key, ""))
            master_repl["{{total_budget}}"] = gio.fmt_currency(grand_total)

            # Replace in template
            doc = Document(str(template_path))
            gio.replace_everywhere(doc, master_repl)

            # Remove unused suffix placeholders
            unused = []
            if not (st.session_state.show_c2 and used_2):
                unused.append("_2")
            if not (st.session_state.show_c3 and used_3):
                unused.append("_3")
            _remove_placeholders_with_suffixes(doc, unused)

            # Save to memory
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            # Filename from Campaign 1
            c1_name = master_repl.get("{{campaign_name}}", "Campaign").replace("/", "-")
            client1 = master_repl.get("{{client_name}}", "Client").replace("/", "-")
            today = dt.date.today().strftime("%Y%m%d")
            outname = f"IO_{client1}_{c1_name}_{today}.docx"

            st.success("IO generated.")
            st.download_button(
                "Download IO",
                data=buf,
                file_name=outname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        st.exception(e)
